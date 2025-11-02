# import docx NOT python-docx
import docx
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime
from matplotlib import font_manager

def list_possible_fonts():
    # Get all font names from system-installed TrueType fonts
    fonts = sorted(set(f.name for f in font_manager.fontManager.ttflist))
    # Print each font name
    for font in fonts:
        print(font)

def docx_example():
    document = Document()
    document.add_paragraph("It was a dark and stormy night.")
    document.save("dark-and-stormy.docx")

    document = Document("dark-and-stormy.docx")
    first_paragraph_text = document.paragraphs[0].text
    print(first_paragraph_text)

class Create_diary:
    def __init__(self, name, font_name, margin_inch):
        dt = datetime.datetime.today()

        self.name = name
        self.date = datetime.datetime(dt.year, dt.month, dt.day)
        self.font_name = font_name
        self.margin_inch = margin_inch

        if name == None:
            self.name = "Audrey Luan"

        if margin_inch == None:
            self.margin_inch = 1

        self.date_title = self.date.strftime("%Y-%m")

        self.split_name = self.name.split(" ")
        self.report_title = f"The {self.split_name[1]} Diaries"
        self.doc = Document()
        self.output_report = f"{self.date_title} diary entries"

        print(f"{self.split_name=}\n{self.date=}\n{self.report_title=}\n{self.output_report=}")

    def set_margins(self):
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

    def edited_by_date(self):
        edited_date = self.date.strftime("%B %Y")
        heading = [self.name, f"Edited in {edited_date}"]
        for item in heading:
            if item is self.name:
                fontsize = 24
            else:
                fontsize = 18
            item_h = self.doc.add_paragraph(item)
            item_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            item_h_format = item_h.runs[0]
            item_h_format.font.size = Pt(fontsize)
            item_h_format.font.name = self.font_name
        return heading

    def add_title(self, user_title=None, fontsize0=12, alignment0=WD_ALIGN_PARAGRAPH.LEFT):
        if user_title is None:
            title = self.doc.add_paragraph(self.report_title)
        else:
            title = self.doc.add_paragraph(user_title)
        title.alignment = alignment0
        title_format = title.runs[0]
        title_format.font.size = Pt(fontsize0)
        title_format.font.name = self.font_name
        return user_title

    def add_paragraph(self, paragraph, fontsize0=12, alignment0=WD_ALIGN_PARAGRAPH.LEFT):
        new_paragraph = self.doc.add_paragraph(paragraph)
        new_paragraph.alignment = alignment0
        new_paragraph_format = new_paragraph.runs[0]
        new_paragraph_format.font.size = Pt(fontsize0)
        new_paragraph_format.font.name = self.font_name
        return paragraph

    def title_page(self, month_date_range):
        self.set_margins()

        title_contents = {None: 24, month_date_range: 18, "\n\n": 18, f"{self.date_title} Diary Entries\nEdition ": 36,
                          "\n\n\n\n": 24}

        for title1, fontsize in title_contents.items():
            self.add_title(
                user_title=title1,
                fontsize0=fontsize,
                alignment0=WD_ALIGN_PARAGRAPH.CENTER
            )
        self.edited_by_date()
        self.doc.add_page_break()
        return title_contents

    def save_diary(self):
        self.doc.save(f"{self.output_report}.docx")
        return f"{self.output_report}.docx"

    def correct_document_pckg(self):
        try:
            document = Document(foo='bar', baz=12)
            print(document['baz'])

        except:
            print("Old Document class is gone — this failed as expected. The docx library is expected.\n")

""" 
if __name__ == '__main__':
    # list_possible_fonts()
    aluan = Create_diary(
        name=None,
        font_name="Courier New",
        margin_inch=None
    )
    aluan.correct_document_pckg()
    
    aluan.title_page()

    aluan.add_paragraph("This is a sample diary entry.") 
"""